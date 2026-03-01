from fastapi import FastAPI, HTTPException, BackgroundTasks, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import os
import json
from pathlib import Path
import logging
import tempfile
import shutil
import asyncio
from typing import Optional, Dict, Any
import uuid
import httpx

from knowledge_base import init_knowledge_base, get_knowledge_base, LinguisticKnowledgeBase
from database_explorer import (
    search_feature_descriptions,
    search_features_by_keywords,
    get_feature_statistics,
    clean_description,
    get_all_feature_ids
)

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 从环境变量获取配置
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT", "8000"))
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173,http://127.0.0.1:5173,http://localhost:3000,http://127.0.0.1:3000").split(",")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE", "50")) * 1024 * 1024  # 默认50MB

app = FastAPI(title="Linguistic Knowledge Base API", version="1.0.0")

# 配置CORS - 使用环境变量
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据模型
class SearchRequest(BaseModel):
    query: str
    n_results: int = 5

class SearchResponse(BaseModel):
    results: List[Dict[str, Any]]
    total_found: int

class KnowledgeBaseInfo(BaseModel):
    total_documents: int
    collection_name: str
    database_path: str
    embedding_method: str
    status: str
    last_updated: Optional[str] = None
    processing_status: Optional[str] = None

class InitRequest(BaseModel):
    openai_api_key: Optional[str] = None

class DocumentRequest(BaseModel):
    file_path: str
    file_type: str  # "pdf" or "csv"

class PaperUploadRequest(BaseModel):
    title: Optional[str] = None
    authors: Optional[str] = None
    abstract: Optional[str] = None
    keywords: Optional[str] = None
    publication_date: Optional[str] = None

class StatusResponse(BaseModel):
    status: str
    message: str
    details: Optional[Dict[str, Any]] = None

# 全局变量
knowledge_base: Optional[LinguisticKnowledgeBase] = None
processing_status: str = "ready"  # idle, processing, completed, error
last_operation: Optional[str] = None
last_operation_time: Optional[str] = None
current_task_id: Optional[str] = None
current_task: Optional[asyncio.Task] = None
task_cancelled: bool = False

@app.on_event("startup")
async def startup_event():
    """应用启动时初始化知识库"""
    global knowledge_base, processing_status
    try:
        # 从环境变量获取API密钥
        openai_api_key = os.getenv("OPENAI_API_KEY")
        knowledge_base = init_knowledge_base(openai_api_key=openai_api_key)
        processing_status = "ready"
        logger.info("知识库初始化完成")
    except Exception as e:
        processing_status = "error"
        logger.error(f"知识库初始化失败: {e}")

@app.post("/api/init", response_model=StatusResponse)
async def initialize_knowledge_base(request: InitRequest):
    """初始化知识库"""
    global knowledge_base
    try:
        knowledge_base = init_knowledge_base(openai_api_key=request.openai_api_key)
        
        return StatusResponse(
            status="success",
            message="知识库初始化成功",
            details={
                "embedding_method": "OpenAI" if request.openai_api_key else "轻量级 TF-IDF",
                "storage_method": "轻量级存储",
                "timestamp": str(Path().cwd())
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"初始化失败: {str(e)}")

@app.get("/api/info", response_model=KnowledgeBaseInfo)
async def get_knowledge_base_info():
    """获取知识库信息"""
    global knowledge_base
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    try:
        info = knowledge_base.get_collection_info()
        
        # 简化状态信息，确保能正确返回
        enhanced_info = {
            **info,
            "status": "ready",
            "last_updated": str(Path().cwd()),
            "processing_status": "ready"
        }
        
        return KnowledgeBaseInfo(**enhanced_info)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取信息失败: {str(e)}")

@app.get("/api/status")
async def get_system_status():
    """获取系统状态"""
    global knowledge_base
    
    try:
        if knowledge_base:
            info = knowledge_base.get_collection_info()
            
            # 只统计 repository 目录下的文件
            repository_dir = Path("../public/repository")
            if repository_dir.exists():
                pdf_files = list(repository_dir.rglob("*.pdf"))
                total_files = len(pdf_files)
            else:
                total_files = 0
            
            return {
                "status": "ready",
                "knowledge_base_initialized": True,
                "total_documents": info.get("total_documents", 0),
                "embedding_method": info.get("embedding_method", "轻量级 TF-IDF"),
                "processing_status": "ready",
                "total_files_in_repository": total_files,
                "pdf_files": total_files,
                "message": f"知识库已初始化，包含 {info.get('total_documents', 0)} 个文档，repository 目录有 {total_files} 个可用PDF文件",
                "database_path": info.get("database_path", "unknown")
            }
        else:
            # 即使未初始化，也统计文件数
            repository_dir = Path("../public/repository")
            if repository_dir.exists():
                pdf_files = list(repository_dir.rglob("*.pdf"))
                total_files = len(pdf_files)
            else:
                total_files = 0
            
            return {
                "status": "not_initialized",
                "knowledge_base_initialized": False,
                "processing_status": "not_initialized",
                "total_files_in_repository": total_files,
                "pdf_files": total_files,
                "message": f"知识库未初始化，repository 目录有 {total_files} 个可用PDF文件"
            }
    except Exception as e:
        return {
            "status": "error",
            "knowledge_base_initialized": False,
            "processing_status": "error",
            "error": str(e)
        }

@app.post("/api/search", response_model=SearchResponse)
async def search_knowledge_base(request: SearchRequest):
    """搜索知识库"""
    global knowledge_base
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    try:
        # 增加默认结果数量，让AI获得更全面的信息
        n_results = max(request.n_results, 10)  # 至少返回10个结果
        results = knowledge_base.search(request.query, n_results)
        return SearchResponse(
            results=results,
            total_found=len(results)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")

@app.post("/api/add-document")
async def add_document(request: DocumentRequest):
    """添加文档到知识库"""
    global knowledge_base
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    try:
        file_path = Path(request.file_path)
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件不存在")
        
        documents = []
        
        if request.file_type.lower() == "pdf":
            # 处理PDF文件
            text = knowledge_base.extract_text_from_pdf(str(file_path))
            if text:
                documents.append({
                    "content": text,
                    "metadata": {
                        "source": str(file_path),
                        "type": "pdf",
                        "filename": file_path.name
                    }
                })
        
        elif request.file_type.lower() == "csv":
            # 处理CSV文件
            documents = knowledge_base.process_csv_data(str(file_path))
        
        else:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        if documents:
            knowledge_base.add_documents(documents)
            return {"message": f"成功添加 {len(documents)} 个文档"}
        else:
            return {"message": "没有找到可处理的文档"}
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"添加文档失败: {str(e)}")

@app.post("/api/upload-paper")
async def upload_paper(
    file: UploadFile = File(...),
    title: Optional[str] = None,
    authors: Optional[str] = None,
    abstract: Optional[str] = None,
    keywords: Optional[str] = None,
    publication_date: Optional[str] = None
):
    """上传论文文件到知识库"""
    global knowledge_base
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    # 检查文件类型
    allowed_types = ["pdf", "docx", "txt"]
    file_extension = file.filename.split(".")[-1].lower() if file.filename else ""
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件类型: {file_extension}。支持的类型: {', '.join(allowed_types)}"
        )
    
            # 检查文件大小限制
        if file.size and file.size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"文件过大: {file.size / 1024 / 1024:.1f}MB。最大允许: {MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )
    
    try:
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            # 写入上传的文件内容
            shutil.copyfileobj(file.file, temp_file)
            temp_file_path = temp_file.name
        
        try:
            documents = []
            extracted_text = ""
            
            if file_extension == "pdf":
                # 处理PDF文件
                text = knowledge_base.extract_text_from_pdf(temp_file_path)
                if text and text.strip():
                    extracted_text = text
                    # 创建元数据，只包含非空值
                    metadata = {
                        "source": file.filename,
                        "type": "paper",
                        "file_type": "pdf",
                        "title": title or file.filename.replace(f".{file_extension}", ""),
                        "upload_time": str(Path().cwd()),
                        "file_size": file.size,
                        "pages": len(text.split('\n')) // 50 + 1  # 估算页数
                    }
                    
                    # 只添加非空的元数据字段
                    if authors:
                        metadata["authors"] = authors
                    if abstract:
                        metadata["abstract"] = abstract
                    if keywords:
                        metadata["keywords"] = keywords
                    if publication_date:
                        metadata["publication_date"] = publication_date
                    
                    documents.append({
                        "content": text,
                        "metadata": metadata
                    })
                else:
                    raise HTTPException(
                        status_code=400,
                        detail="PDF文件内容为空或无法提取文本"
                    )
            
            elif file_extension == "docx":
                # 处理Word文档
                try:
                    import docx
                    doc = docx.Document(temp_file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    if text.strip():
                        extracted_text = text
                        # 创建元数据，只包含非空值
                        metadata = {
                            "source": file.filename,
                            "type": "paper",
                            "file_type": "docx",
                            "title": title or file.filename.replace(f".{file_extension}", ""),
                            "upload_time": str(Path().cwd()),
                            "file_size": file.size,
                            "paragraphs": len(doc.paragraphs)
                        }
                        
                        # 只添加非空的元数据字段
                        if authors:
                            metadata["authors"] = authors
                        if abstract:
                            metadata["abstract"] = abstract
                        if keywords:
                            metadata["keywords"] = keywords
                        if publication_date:
                            metadata["publication_date"] = publication_date
                        
                        documents.append({
                            "content": text,
                            "metadata": metadata
                        })
                    else:
                        raise HTTPException(
                            status_code=400,
                            detail="Word文档内容为空"
                        )
                except ImportError:
                    raise HTTPException(
                        status_code=500, 
                        detail="需要安装python-docx包来处理Word文档"
                    )
            
            elif file_extension == "txt":
                # 处理文本文件
                try:
                    with open(temp_file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    if text.strip():
                        extracted_text = text
                        # 创建元数据，只包含非空值
                        metadata = {
                            "source": file.filename,
                            "type": "paper",
                            "file_type": "txt",
                            "title": title or file.filename.replace(f".{file_extension}", ""),
                            "upload_time": str(Path().cwd()),
                            "file_size": file.size,
                            "characters": len(text),
                            "lines": len(text.split('\n'))
                        }
                        
                        # 只添加非空的元数据字段
                        if authors:
                            metadata["authors"] = authors
                        if abstract:
                            metadata["abstract"] = abstract
                        if keywords:
                            metadata["keywords"] = keywords
                        if publication_date:
                            metadata["publication_date"] = publication_date
                        
                        documents.append({
                            "content": text,
                            "metadata": metadata
                        })
                    else:
                        raise HTTPException(
                            status_code=400,
                            detail="文本文件内容为空"
                        )
                except UnicodeDecodeError:
                    # 尝试其他编码
                    try:
                        with open(temp_file_path, 'r', encoding='gbk') as f:
                            text = f.read()
                        if text.strip():
                            extracted_text = text
                            # 创建元数据，只包含非空值
                            metadata = {
                                "source": file.filename,
                                "type": "paper",
                                "file_type": "txt",
                                "title": title or file.filename.replace(f".{file_extension}", ""),
                                "upload_time": str(Path().cwd()),
                                "file_size": file.size,
                                "characters": len(text),
                                "lines": len(text.split('\n')),
                                "encoding": "gbk"
                            }
                            
                            # 只添加非空的元数据字段
                            if authors:
                                metadata["authors"] = authors
                            if abstract:
                                metadata["abstract"] = abstract
                            if keywords:
                                metadata["keywords"] = keywords
                            if publication_date:
                                metadata["publication_date"] = publication_date
                            
                            documents.append({
                                "content": text,
                                "metadata": metadata
                            })
                        else:
                            raise HTTPException(
                                status_code=400,
                                detail="文本文件内容为空"
                            )
                    except UnicodeDecodeError:
                        raise HTTPException(
                            status_code=400,
                            detail="无法解码文本文件，请检查文件编码"
                        )
            
            if documents:
                # 清理元数据，移除None值
                for doc in documents:
                    if "metadata" in doc:
                        # 过滤掉值为None的元数据字段
                        doc["metadata"] = {k: v for k, v in doc["metadata"].items() if v is not None}
                
                # 保存文件到repository目录
                repository_dir = Path("../public/repository")
                repository_dir.mkdir(exist_ok=True)
                
                # 生成唯一的文件名（避免重名）
                base_name = Path(file.filename).stem
                counter = 1
                final_filename = file.filename
                
                while (repository_dir / final_filename).exists():
                    name_parts = Path(file.filename).stem, Path(file.filename).suffix
                    final_filename = f"{name_parts[0]}_{counter}{name_parts[1]}"
                    counter += 1
                
                final_file_path = repository_dir / final_filename
                
                # 复制文件到repository目录
                shutil.copy2(temp_file_path, final_file_path)
                logger.info(f"文件已保存到repository目录: {final_file_path}")
                
                # 更新元数据中的source路径
                for doc in documents:
                    doc["metadata"]["source"] = str(final_file_path)
                    doc["metadata"]["repository_filename"] = final_filename
                
                # 添加到知识库
                knowledge_base.add_documents(documents)
                
                # 获取文档统计信息
                doc_info = knowledge_base.get_collection_info()
                
                return {
                    "message": f"论文上传成功，已保存到repository目录",
                    "filename": file.filename,
                    "repository_filename": final_filename,
                    "file_type": file_extension,
                    "documents_added": len(documents),
                    "file_size": file.size,
                    "text_length": len(extracted_text),
                    "collection_info": {
                        "total_documents": doc_info.get("total_documents", 0),
                        "collection_name": doc_info.get("collection_name", ""),
                        "database_path": doc_info.get("database_path", "")
                    },
                    "metadata": {
                        "title": documents[0]["metadata"]["title"],
                        "authors": authors,
                        "abstract": abstract,
                        "keywords": keywords,
                        "publication_date": publication_date
                    }
                }
            else:
                raise HTTPException(
                    status_code=400,
                    detail="文件内容为空，无法添加到知识库"
                )
                
        finally:
            # 清理临时文件
            try:
                os.unlink(temp_file_path)
            except OSError:
                logger.warning(f"无法删除临时文件: {temp_file_path}")
            
    except HTTPException:
        # 重新抛出HTTP异常
        raise
    except Exception as e:
        logger.error(f"论文上传失败: {e}")
        raise HTTPException(
            status_code=500, 
            detail=f"论文上传失败: {str(e)}"
        )

@app.post("/api/add-documents-batch")
async def add_documents_batch(background_tasks: BackgroundTasks):
    """批量添加文档（后台任务）"""
    global knowledge_base, current_task_id, current_task, task_cancelled
    
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    # 如果已有任务在运行，返回错误
    if current_task and not current_task.done():
        return {
            "status": "error",
            "message": "已有任务正在运行，请等待完成或先取消当前任务",
            "current_task_id": current_task_id
        }
    
    # 重置取消标志
    task_cancelled = False
    
    async def process_documents():
        """异步处理文档的函数"""
        global current_task_id, task_cancelled
        
        try:
            # 只处理 repository 目录下的 Handbook1 和 Handbook2
            repository_dir = Path("../public/repository")
            handbook_dirs = ["Handbook1", "Handbook2"]
            pdf_files = []

            if repository_dir.exists():
                # 本地路径存在，直接使用
                for handbook_dir_name in handbook_dirs:
                    handbook_dir = repository_dir / handbook_dir_name
                    if handbook_dir.exists():
                        handbook_pdfs = list(handbook_dir.rglob("*.pdf"))
                        pdf_files.extend(handbook_pdfs)
                        logger.info(f"找到 {handbook_dir_name} 目录: {len(handbook_pdfs)} 个PDF文件")
                    else:
                        logger.warning(f"{handbook_dir_name} 目录不存在")
            else:
                # 本地路径不存在（HF Space），从 GitHub 下载 PDF
                logger.info("本地 repository 目录不存在，从 GitHub 下载 PDF...")
                GITHUB_API_BASE = "https://api.github.com/repos/yutong-yang/linguistic-data-visualization/contents/public/repository"
                temp_dir = Path("/tmp/handbook_cache")
                temp_dir.mkdir(exist_ok=True)

                for handbook_name in handbook_dirs:
                    if task_cancelled:
                        break
                    try:
                        api_url = f"{GITHUB_API_BASE}/{handbook_name}"
                        api_resp = await asyncio.get_event_loop().run_in_executor(
                            None,
                            lambda url=api_url: httpx.get(url, timeout=30, headers={"Accept": "application/vnd.github+json"})
                        )
                        if api_resp.status_code != 200:
                            logger.error(f"GitHub API 请求失败 {handbook_name}: {api_resp.status_code}")
                            continue

                        file_entries = [f for f in api_resp.json() if f.get("name", "").lower().endswith(".pdf") and f.get("type") == "file"]
                        logger.info(f"GitHub {handbook_name}: 找到 {len(file_entries)} 个PDF")

                        handbook_temp_dir = temp_dir / handbook_name
                        handbook_temp_dir.mkdir(exist_ok=True)

                        for file_info in file_entries:
                            if task_cancelled:
                                break
                            filename = file_info["name"]
                            download_url = file_info["download_url"]
                            local_path = handbook_temp_dir / filename

                            if not local_path.exists():
                                logger.info(f"下载: {filename}")
                                pdf_resp = await asyncio.get_event_loop().run_in_executor(
                                    None,
                                    lambda url=download_url: httpx.get(url, timeout=120, follow_redirects=True)
                                )
                                if pdf_resp.status_code == 200:
                                    local_path.write_bytes(pdf_resp.content)
                                    logger.info(f"下载完成: {filename}")
                                else:
                                    logger.error(f"下载失败: {filename}, 状态码: {pdf_resp.status_code}")
                                    continue
                            else:
                                logger.info(f"已缓存，跳过下载: {filename}")

                            pdf_files.append(local_path)

                    except Exception as e:
                        logger.error(f"从 GitHub 获取 {handbook_name} 失败: {e}")
            
            total_files = len(pdf_files)
            
            if total_files == 0:
                logger.info("Handbook1 和 Handbook2 目录中没有PDF文件")
                return
            
            logger.info(f"开始批量处理 Handbook1 和 Handbook2 目录，总计 {total_files} 个PDF文件")
            
            processed_files = 0
            errors = []
            added_files = 0
            
            # 获取已存在的文档源文件列表
            existing_sources = knowledge_base.collection.get_existing_sources()
            logger.info(f"知识库中已有 {len(existing_sources)} 个文档源")
            
            # 处理PDF文件
            for pdf_file in pdf_files:
                # 检查是否被取消
                if task_cancelled:
                    logger.info("任务被用户取消")
                    break
                
                try:
                    # 检查文件是否已经处理过
                    file_path_str = str(pdf_file)
                    if file_path_str in existing_sources:
                        logger.info(f"文件已存在，跳过: {pdf_file.name}")
                        processed_files += 1
                        continue
                    
                    text = knowledge_base.extract_text_from_pdf(file_path_str)
                    if text:
                        filename_base = pdf_file.stem  # 不含扩展名的文件名
                        is_doi_format = bool(filename_base and len(filename_base) > 10 and '-' in filename_base and any(c.isdigit() for c in filename_base))
                        
                        # 确定使用的标题（优先使用文件名，因为文件名通常就是论文标题）
                        if is_doi_format:
                            # 如果文件名是DOI格式（无意义的ID），尝试从PDF中提取标题
                            extracted_title = knowledge_base.extract_title_from_pdf(file_path_str)
                            if extracted_title:
                                doc_title = extracted_title
                                logger.info(f"  - DOI格式文件名，使用提取的标题: {doc_title}")
                            else:
                                # 提取失败，使用文件名
                                doc_title = filename_base
                                logger.warning(f"  - DOI格式文件名，无法提取标题，使用文件名: {doc_title}")
                        else:
                            # 文件名不是DOI格式，直接使用文件名作为标题（文件名通常就是论文标题）
                            doc_title = filename_base
                            logger.info(f"  - 使用文件名作为标题: {doc_title}")
                            
                            # 可选：验证文件名是否看起来像标题，如果不是，尝试从PDF提取
                            # 如果文件名太短或包含特殊字符，可能是无意义的ID
                            if len(filename_base) < 10 or filename_base.replace('_', '').replace('-', '').isdigit():
                                extracted_title = knowledge_base.extract_title_from_pdf(file_path_str)
                                if extracted_title and len(extracted_title) > len(filename_base):
                                    doc_title = extracted_title
                                    logger.info(f"  - 文件名看起来不像标题，使用提取的标题: {doc_title}")
                        
                        # 确定文件来自哪个 handbook 目录
                        handbook_name = None
                        if "Handbook1" in file_path_str:
                            handbook_name = "Handbook1"
                        elif "Handbook2" in file_path_str:
                            handbook_name = "Handbook2"
                        
                        documents = [{
                            "content": text,
                            "metadata": {
                                "source": file_path_str,
                                "type": "pdf",
                                "filename": pdf_file.name,
                                "title": doc_title,  # 添加提取的标题
                                "handbook": handbook_name,  # 标记来源目录
                                "processed_time": str(Path().cwd())
                            }
                        }]
                        knowledge_base.add_documents(documents)
                        processed_files += 1
                        added_files += 1
                        existing_sources.add(file_path_str)  # 添加到已处理列表
                        logger.info(f"成功处理PDF文件: {pdf_file.name} ({processed_files}/{total_files})")
                    else:
                        logger.warning(f"PDF文件内容为空: {pdf_file.name}")
                        processed_files += 1
                        
                except Exception as e:
                    error_msg = f"处理PDF文件失败 {pdf_file}: {e}"
                    errors.append(error_msg)
                    logger.error(error_msg)
                    processed_files += 1  # 即使失败也计数
            
            if task_cancelled:
                logger.info(f"任务被取消，已处理 {processed_files}/{total_files} 个文件")
            else:
                logger.info(f"批量文档处理完成: 总计 {total_files} 个文件，成功处理 {processed_files} 个，新增 {added_files} 个，错误 {len(errors)} 个")
            
        except Exception as e:
            logger.error(f"批量处理文档失败: {e}")
        finally:
            # 清理任务状态
            current_task = None
            current_task_id = None
    
    # 创建新任务
    current_task_id = str(uuid.uuid4())
    current_task = asyncio.create_task(process_documents())
    
    # 立即返回，包含文件统计信息（只统计 Handbook1 和 Handbook2）
    repository_dir = Path("../public/repository")
    total_files = 0
    if repository_dir.exists():
        handbook_dirs = ["Handbook1", "Handbook2"]
        pdf_files = []
        for handbook_dir_name in handbook_dirs:
            handbook_dir = repository_dir / handbook_dir_name
            if handbook_dir.exists():
                handbook_pdfs = list(handbook_dir.rglob("*.pdf"))
                pdf_files.extend(handbook_pdfs)
        total_files = len(pdf_files)
    else:
        total_files = -1  # 表示将从 GitHub 下载，数量待定
    
    return {
        "status": "processing",
        "message": f"批量处理已启动",
        "total_files": total_files,
        "processed_files": 0,
        "operation": "batch_processing",
        "task_id": current_task_id
    }

@app.post("/api/cancel-task")
async def cancel_current_task():
    """取消当前正在运行的任务"""
    global current_task, task_cancelled, current_task_id
    
    if not current_task or current_task.done():
        return {
            "status": "error",
            "message": "没有正在运行的任务"
        }
    
    try:
        # 设置取消标志
        task_cancelled = True
        
        # 取消任务
        current_task.cancel()
        
        # 等待任务完成
        try:
            await current_task
        except asyncio.CancelledError:
            pass
        
        # 清理状态
        current_task = None
        current_task_id = None
        task_cancelled = False
        
        return {
            "status": "success",
            "message": "任务已成功取消",
            "task_id": current_task_id
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"取消任务失败: {str(e)}"
        }

@app.get("/api/task-status")
async def get_task_status():
    """获取当前任务状态"""
    global current_task, current_task_id, task_cancelled
    
    if not current_task:
        return {
            "status": "idle",
            "message": "没有正在运行的任务",
            "task_id": None
        }
    
    if current_task.done():
        if current_task.cancelled():
            status = "cancelled"
            message = "任务已被取消"
        elif current_task.exception():
            status = "error"
            message = f"任务执行出错: {current_task.exception()}"
        else:
            status = "completed"
            message = "任务已完成"
        
        # 清理已完成的任务
        current_task = None
        current_task_id = None
        task_cancelled = False
        
        return {
            "status": status,
            "message": message,
            "task_id": current_task_id
        }
    
    return {
        "status": "running",
        "message": "任务正在运行中",
        "task_id": current_task_id,
        "cancellable": True
    }

@app.delete("/api/clear")
async def clear_knowledge_base():
    """清空知识库"""
    global knowledge_base, processing_status, last_operation, last_operation_time
    if not knowledge_base:
        raise HTTPException(status_code=400, detail="知识库未初始化")
    
    try:
        processing_status = "processing"
        last_operation = "clearing"
        last_operation_time = str(Path().cwd())
        
        knowledge_base.clear_collection()
        
        processing_status = "ready"
        last_operation = "clearing_completed"
        last_operation_time = str(Path().cwd())
        
        return StatusResponse(
            status="success",
            message="知识库已清空",
            details={
                "operation": "clear",
                "timestamp": last_operation_time
            }
        )
    except Exception as e:
        processing_status = "error"
        last_operation = "clearing_error"
        last_operation_time = str(Path().cwd())
        raise HTTPException(status_code=500, detail=f"清空失败: {str(e)}")

@app.get("/api/health")
async def health_check():
    """健康检查"""
    global knowledge_base
    return {
        "status": "healthy", 
        "knowledge_base_initialized": knowledge_base is not None,
        "processing_status": "ready" if knowledge_base else "not_initialized",
        "timestamp": str(Path().cwd())
    }

@app.get("/api/progress")
async def get_processing_progress():
    """获取处理进度"""
    try:
        # 检查是否有正在进行的批量处理
        # 这里可以通过检查日志或临时文件来判断进度
        public_dir = Path("../public")
        pdf_files = list(public_dir.rglob("*.pdf"))
        csv_files = list(public_dir.rglob("*.csv"))
        total_files = len(pdf_files) + len(csv_files)
        
        # 获取当前知识库中的文档数
        if knowledge_base:
            current_docs = knowledge_base.get_collection_info().get("total_documents", 0)
        else:
            current_docs = 0
        
        return {
            "status": "processing" if total_files > 0 else "idle",
            "total_files": total_files,
            "current_documents": current_docs,
            "message": f"总计 {total_files} 个文件，当前知识库有 {current_docs} 个文档"
        }
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }

# 千问API代理端点
class QianwenRequest(BaseModel):
    prompt: str
    model: Optional[str] = "qwen-turbo"
    temperature: Optional[float] = 0.7
    max_tokens: Optional[int] = 2000

class GeminiRequest(BaseModel):
    prompt: str
    model: Optional[str] = "gemini-1.5-flash"
    temperature: Optional[float] = 0.7
    max_tokens: Optional[int] = 2000

class FeatureRecommendationRequest(BaseModel):
    user_query: str
    language_data: Optional[List[Dict[str, Any]]] = None
    feature_descriptions: Optional[Dict[str, Any]] = None
    n_kb_results: Optional[int] = 10
    api_provider: Optional[str] = "gemini"  # "gemini" 或 "qianwen"
    lang: Optional[str] = "en"  # 用户界面语言 "en" 或 "zh"

@app.post("/api/gemini/chat")
async def gemini_chat(
    request: GeminiRequest,
    http_request: Request
):
    """Gemini API代理端点，用于避免CORS问题和保护API Key"""
    try:
        # 优先从环境变量获取API密钥
        api_key = os.getenv("GEMINI_API_KEY")
        
        # 如果没有环境变量，尝试从请求头获取
        if not api_key:
            api_key = http_request.headers.get("X-API-Key") or http_request.headers.get("Authorization", "").replace("Bearer ", "")
        
        # 如果仍然没有API Key，返回错误
        if not api_key:
            raise HTTPException(
                status_code=400, 
                detail="GEMINI_API_KEY未设置。请在环境变量中设置GEMINI_API_KEY，或在请求头中提供X-API-Key"
            )
        
        # Gemini API端点
        api_url = "https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent"
        
        # 构建请求体
        payload = {
            "contents": [{
                "parts": [{
                    "text": request.prompt
                }]
            }]
        }
        
        # 使用httpx异步调用Gemini API
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{api_url}?key={api_key}",
                headers={
                    "Content-Type": "application/json"
                },
                json=payload
            )
            
            if response.status_code != 200:
                error_text = response.text
                logger.error(f"Gemini API error: {response.status_code} - {error_text}")
                raise HTTPException(
                    status_code=response.status_code,
                    detail=f"Gemini API调用失败: {error_text}"
                )
            
            try:
                data = response.json()
            except Exception as e:
                logger.error(f"Failed to parse Gemini API response as JSON: {e}. Response text: {response.text}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Gemini API返回的不是有效的JSON格式: {str(e)}"
                )
            
            # 提取响应内容
            if data.get("candidates") and len(data["candidates"]) > 0:
                candidate = data["candidates"][0]
                if candidate.get("content") and candidate["content"].get("parts"):
                    content = candidate["content"]["parts"][0].get("text", "")
                    if content:
                        return {
                            "success": True,
                            "content": str(content),
                            "usage": data.get("usage", {})
                        }
            
            # 如果所有格式都不匹配，记录完整的响应以便调试
            response_str = json.dumps(data, ensure_ascii=False, indent=2) if isinstance(data, (dict, list)) else str(data)
            logger.error(f"Unexpected Gemini API response structure. Full response: {response_str}")
            raise HTTPException(
                status_code=500,
                detail=f"Gemini API返回格式异常。无法从响应中提取内容。响应结构: {response_str}"
            )
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="请求超时，请稍后重试")
    except httpx.RequestError as e:
        logger.error(f"Gemini API request error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求失败: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"调用Gemini API时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调用Gemini API时出错: {str(e)}")

@app.post("/api/qianwen/chat")
async def qianwen_chat(
    request: QianwenRequest,
    http_request: Request
):
    """千问API代理端点，用于避免CORS问题"""
    try:
        # 优先从环境变量获取API密钥
        api_key = os.getenv("QIANWEN_API_KEY")
        
        # 如果没有环境变量，尝试从请求头获取
        if not api_key:
            api_key = http_request.headers.get("X-API-Key") or http_request.headers.get("Authorization", "").replace("Bearer ", "")
        
        # 如果仍然没有API Key，返回错误
        if not api_key:
            raise HTTPException(
                status_code=400, 
                detail="QIANWEN_API_KEY未设置。请在环境变量中设置QIANWEN_API_KEY，或在启动后端时设置：export QIANWEN_API_KEY=your_key"
            )
        
        # 千问API端点
        api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
        # 构建请求体
        payload = {
            "model": request.model,
            "input": {
                "messages": [
                    {
                        "role": "user",
                        "content": request.prompt
                    }
                ]
            },
            "parameters": {
                "temperature": request.temperature,
                "max_tokens": request.max_tokens
            }
        }
        
        # 使用httpx异步调用千问API
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                api_url,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                },
                json=payload
            )
            
            if response.status_code != 200:
                error_text = response.text
                logger.error(f"Qianwen API error: {response.status_code} - {error_text}")
                raise HTTPException(
                    status_code=response.status_code,
                    detail=f"千问API调用失败: {error_text}"
                )
            
            try:
                data = response.json()
            except Exception as e:
                logger.error(f"Failed to parse Qianwen API response as JSON: {e}. Response text: {response.text}")
                raise HTTPException(
                    status_code=500,
                    detail=f"千问API返回的不是有效的JSON格式: {str(e)}"
                )
            
            # 记录响应结构以便调试（只记录前500字符）
            logger.info(f"Qianwen API response keys: {list(data.keys()) if isinstance(data, dict) else 'Not a dict'}")
            
            # 提取响应内容 - 尝试多种可能的响应格式
            content = None
            
            # 格式1: output.text (千问API的标准格式)
            if isinstance(data, dict) and data.get("output"):
                output = data["output"]
                if isinstance(output, dict):
                    # 千问API返回格式: {"output": {"text": "...", "finish_reason": "stop"}}
                    content = output.get("text")
                    # 如果没有text，尝试choices格式（兼容其他可能的格式）
                    if not content and output.get("choices"):
                        choices = output["choices"]
                        if isinstance(choices, list) and len(choices) > 0:
                            choice = choices[0]
                            if isinstance(choice, dict):
                                # 尝试获取message.content
                                message = choice.get("message", {})
                                if isinstance(message, dict):
                                    content = message.get("content", "")
                                elif isinstance(message, str):
                                    content = message
                                # 如果message不存在，尝试直接获取content或text
                                if not content:
                                    content = choice.get("content") or choice.get("text")
            
            # 格式2: 直接是content字段
            if not content and isinstance(data, dict) and data.get("content"):
                content = data["content"]
            
            # 格式3: result字段
            if not content and isinstance(data, dict) and data.get("result"):
                content = data["result"]
            
            # 格式4: text字段（顶层）
            if not content and isinstance(data, dict) and data.get("text"):
                content = data["text"]
            
            # 格式5: 如果data本身就是字符串
            if not content and isinstance(data, str):
                content = data
            
            if content:
                return {
                    "success": True,
                    "content": str(content),
                    "usage": data.get("usage", {}) if isinstance(data, dict) else {}
                }
            else:
                # 如果所有格式都不匹配，记录完整的响应以便调试
                response_str = json.dumps(data, ensure_ascii=False, indent=2) if isinstance(data, (dict, list)) else str(data)
                logger.error(f"Unexpected Qianwen API response structure. Full response: {response_str}")
                raise HTTPException(
                    status_code=500,
                    detail=f"千问API返回格式异常。无法从响应中提取内容。响应结构: {response_str}"
                )
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="请求超时，请稍后重试")
    except httpx.RequestError as e:
        logger.error(f"Qianwen API request error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"网络请求失败: {str(e)}")
    except Exception as e:
        logger.error(f"Qianwen API error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调用千问API时出错: {str(e)}")

async def call_qianwen_api_for_recommendation(prompt: str, api_key: str) -> str:
    """调用千问API用于特征推荐"""
    try:
        api_url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        
        payload = {
            "model": "qwen-turbo",
            "input": {
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            },
            "parameters": {
                "temperature": 0.7,
                "max_tokens": 2000
            }
        }
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                api_url,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                },
                json=payload
            )
            
            if response.status_code != 200:
                raise HTTPException(
                    status_code=response.status_code,
                    detail=f"千问API 请求失败: {response.text}"
                )
            
            data = response.json()
            
            # 提取响应内容
            if isinstance(data, dict) and data.get("output"):
                output = data["output"]
                if isinstance(output, dict):
                    content = output.get("text")
                    if content:
                        return str(content)
            
            raise HTTPException(status_code=500, detail="千问API 返回格式异常")
            
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="请求超时，请稍后重试")
    except Exception as e:
        logger.error(f"调用千问API时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调用千问API时出错: {str(e)}")

async def call_gemini_api(prompt: str, api_key: str) -> str:
    """调用 Gemini API"""
    try:
        api_url = "https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent"
        
        payload = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }]
        }
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{api_url}?key={api_key}",
                headers={"Content-Type": "application/json"},
                json=payload
            )
            
            if response.status_code != 200:
                raise HTTPException(
                    status_code=response.status_code,
                    detail=f"Gemini API 请求失败: {response.text}"
                )
            
            data = response.json()
            
            if data.get("candidates") and len(data["candidates"]) > 0:
                candidate = data["candidates"][0]
                if candidate.get("content") and candidate["content"].get("parts"):
                    return candidate["content"]["parts"][0].get("text", "")
            
            raise HTTPException(status_code=500, detail="Gemini API 返回格式异常")
            
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="请求超时，请稍后重试")
    except Exception as e:
        logger.error(f"调用 Gemini API 时出错: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调用 Gemini API 时出错: {str(e)}")

def extract_features_from_knowledge_base(kb_results: List[Dict[str, Any]]) -> List[str]:
    """从知识库搜索结果中提取特征 ID"""
    features = set()
    import re
    
    for result in kb_results:
        document = result.get("document", {})
        content = (result.get("content") or document.get("content") or "").lower()
        
        # 提取 GB 特征
        gb_matches = re.findall(r'gb\d{3}', content, re.IGNORECASE)
        for match in gb_matches:
            feature_id = match.upper()
            num = int(feature_id[2:])
            if 20 <= num <= 999:
                features.add(feature_id)
        
        # 提取 EA 特征
        ea_matches = re.findall(r'ea\d{3}', content, re.IGNORECASE)
        for match in ea_matches:
            features.add(match.upper())
        
        # 提取环境特征
        if "richness" in content:
            features.update([
                "AmphibianRichness",
                "BirdRichness",
                "MammalRichness",
                "VascularPlantsRichness"
            ])
    
    return list(features)

@app.post("/api/feature-recommendation")
async def recommend_features(
    request: FeatureRecommendationRequest,
    http_request: Request
):
    """特征推荐端点"""
    try:
        logger.info("=" * 60)
        logger.info("🔍 开始特征推荐流程")
        logger.info(f"用户查询: {request.user_query}")
        logger.info(f"API Provider: {request.api_provider or 'gemini'}")
        logger.info(f"界面语言: {request.lang or 'en'}")
        
        # 确定使用的 API provider
        api_provider = request.api_provider or "gemini"
        user_lang = request.lang or "en"
        
        # 获取 API Key
        api_key = None
        if api_provider == "qianwen":
            # 优先从环境变量获取
            api_key = os.getenv("QIANWEN_API_KEY")
            # 如果没有环境变量，从请求头获取
            if not api_key:
                api_key = http_request.headers.get("X-API-Key") or http_request.headers.get("Authorization", "").replace("Bearer ", "")
            
            if not api_key:
                raise HTTPException(
                    status_code=400,
                    detail="QIANWEN_API_KEY未设置。请在环境变量中设置QIANWEN_API_KEY，或在请求头中提供X-API-Key"
                )
        else:
            # Gemini
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                api_key = http_request.headers.get("X-API-Key") or http_request.headers.get("Authorization", "").replace("Bearer ", "")
            
            if not api_key:
                raise HTTPException(
                    status_code=400,
                    detail="GEMINI_API_KEY未设置。请在环境变量中设置GEMINI_API_KEY，或在请求头中提供X-API-Key"
                )
        
        # 1. 获取数据库统计信息
        logger.info("📊 步骤1: 获取数据库统计信息")
        stats = get_feature_statistics()
        logger.info(f"  - Grambank特征数: {stats.get('totalGrambankFeatures', 0)}")
        logger.info(f"  - D-PLACE特征数: {stats.get('totalDplaceFeatures', 0)}")
        logger.info(f"  - WALS特征数: {stats.get('totalWalsFeatures', 0)}")
        
        # 2. 先进行简单的关键词匹配（优先展示直接包含关键词的特征）
        logger.info("🔑 步骤2a: 关键词直接匹配（简单查找）")
        keyword_matches = search_features_by_keywords(request.user_query, 30)
        logger.info(f"  - 关键词匹配找到 {len(keyword_matches)} 个特征")
        if len(keyword_matches) > 0:
            logger.info(f"  - 前10个关键词匹配: {[f.get('id', '') for f in keyword_matches[:10]]}")
        
        # 2b. 搜索相关特征描述（CSV数据库）- 使用自然语言搜索，返回更多结果让LLM自己判断
        logger.info("🔎 步骤2b: 搜索CSV数据库中的相关特征（自然语言搜索）")
        semantic_search_results = search_feature_descriptions(request.user_query, 50)  # 增加搜索数量，让LLM有更多选择
        logger.info(f"  - 语义搜索找到 {len(semantic_search_results)} 个相关特征")
        if len(semantic_search_results) > 0:
            logger.info(f"  - 前10个语义搜索结果: {[f.get('id', '') for f in semantic_search_results[:10]]}")
        
        # 合并结果：关键词匹配的特征优先，然后添加语义搜索结果中不重复的特征
        keyword_match_ids = {f.get('id') for f in keyword_matches}
        search_results = keyword_matches.copy()  # 先添加关键词匹配的结果
        for feature in semantic_search_results:
            if feature.get('id') not in keyword_match_ids:
                search_results.append(feature)
        
        logger.info(f"  - 合并后总共 {len(search_results)} 个特征（{len(keyword_matches)} 个关键词匹配 + {len(semantic_search_results) - len([f for f in semantic_search_results if f.get('id') in keyword_match_ids])} 个语义搜索）")
        
        # 3. 搜索知识库（PDF文档）
        logger.info("📚 步骤3: 搜索知识库（PDF文献）")
        knowledge_base_context = ""
        extracted_features = []
        
        if knowledge_base:
            try:
                kb_results = knowledge_base.search(request.user_query, request.n_kb_results or 10)
                logger.info(f"  - 知识库搜索结果数: {len(kb_results) if kb_results else 0}")
                
                if kb_results and len(kb_results) > 0:
                    kb_title = "知识库（PDF研究文献）相关内容" if user_lang == "zh" else "Knowledge Base (PDF Research Literature) Related Content"
                    knowledge_base_context = f"\n=== {kb_title} ===\n"
                    # 存储文档映射，用于后续引用
                    document_mapping = {}
                    
                    # 第一步：先收集所有文档的标题，构建文档名称列表
                    for index, result in enumerate(kb_results[:10], 1):
                        document = result.get("document", {})
                        metadata = result.get("metadata") or document.get("metadata", {})
                        
                        # 完全依赖metadata中的title（文档添加时已提取并保存）
                        doc_title = metadata.get("title")
                        
                        # 如果metadata中有title，直接使用（这是最可靠的）
                        if doc_title and len(doc_title.strip()) >= 5:
                            doc_title = doc_title.strip()
                            # 排除明显的内容片段特征
                            if (not doc_title.endswith('?') and 
                                not doc_title.startswith('According to') and
                                not doc_title.startswith('Document')):
                                logger.debug(f"  - 文档 {index} 使用metadata中的title: {doc_title}")
                                pass  # 使用metadata中的标题
                            else:
                                logger.warning(f"  - 文档 {index} metadata中的title看起来不合理，跳过: {doc_title}")
                                doc_title = None
                        
                        # 如果metadata中没有title，使用文件名作为fallback
                        if not doc_title or len(doc_title.strip()) < 5:
                            doc_title = metadata.get("filename") or metadata.get("source") or ("未知文档" if user_lang == "zh" else "Unknown Document")
                            # 提取文件名（去除路径）
                            if "/" in doc_title:
                                doc_title = doc_title.split("/")[-1]
                            if "\\" in doc_title:
                                doc_title = doc_title.split("\\")[-1]
                            # 去除文件扩展名
                            if doc_title.endswith(('.pdf', '.PDF', '.docx', '.DOCX', '.txt', '.TXT')):
                                doc_title = doc_title.rsplit('.', 1)[0]
                            
                            if '-' in doc_title and len(doc_title) > 15 and any(c.isdigit() for c in doc_title):
                                logger.warning(f"  - 文档 {index} 使用DOI格式文件名作为标题: {doc_title} (metadata中缺少title，建议重新添加文档)")
                            else:
                                logger.info(f"  - 文档 {index} 使用文件名作为标题: {doc_title} (metadata中缺少title)")
                        
                        # 确保doc_title不为空
                        if not doc_title or len(doc_title.strip()) < 3:
                            doc_title = f"文档 {index}" if user_lang == "zh" else f"Document {index}"
                            logger.warning(f"  - 文档 {index} 无法获取标题，使用默认名称")
                        
                        document_mapping[index] = doc_title
                    
                    # 第二步：在展示内容之前，先展示文档名称列表和引用说明（重要！）
                    if document_mapping:
                        if user_lang == "zh":
                            doc_list = "\n".join([f"  - 文献{idx}: {name}" for idx, name in document_mapping.items()])
                            knowledge_base_context += f"\n**⚠️ 重要：文档名称列表（引用时必须使用这些确切名称）**：\n{doc_list}\n\n"
                            knowledge_base_context += f"**引用格式要求（必须严格遵守）**：\n"
                            knowledge_base_context += f"1. 引用文献时，直接使用文档名称，格式为：\"根据{list(document_mapping.values())[0] if document_mapping else '[文档名称]'}\" 或 \"{list(document_mapping.values())[1] if len(document_mapping) > 1 else '[文档名称]'}提到...\"\n"
                            knowledge_base_context += f"2. **绝对禁止**：使用\"文献1\"、\"Document 1\"等编号前缀\n"
                            knowledge_base_context += f"3. **绝对禁止**：使用PDF内容片段作为文档名称（如\"Studies in Language Companion Series\"或\"(Studies in Language Companion Series) Matti M\"是错误的）\n"
                            knowledge_base_context += f"4. **必须**：使用文档名称列表中列出的确切名称，不要修改、截断或从内容中提取\n"
                            knowledge_base_context += f"5. 正确示例：\"根据{list(document_mapping.values())[0] if document_mapping else '[文档名称]'}\" ✅\n"
                            knowledge_base_context += f"6. 错误示例：\"根据文献1: {list(document_mapping.values())[0] if document_mapping else '[文档名称]'}\" ❌（不要使用编号）\n"
                            knowledge_base_context += f"7. 错误示例：\"Document 1: Studies in Language Companion Series\" ❌（这是内容片段，不是文档名称）\n"
                        else:
                            doc_list = "\n".join([f"  - Document {idx}: {name}" for idx, name in document_mapping.items()])
                            knowledge_base_context += f"\n**⚠️ IMPORTANT: Document Name List (MUST use these exact names when citing)**：\n{doc_list}\n\n"
                            knowledge_base_context += f"**Citation Format Requirements (MUST strictly follow)**：\n"
                            knowledge_base_context += f"1. When citing documents, directly use the document name, format: \"According to {list(document_mapping.values())[0] if document_mapping else '[Document Name]'}\" or \"{list(document_mapping.values())[1] if len(document_mapping) > 1 else '[Document Name]'} mentions...\"\n"
                            knowledge_base_context += f"2. **ABSOLUTELY FORBIDDEN**: Using \"Document 1\", \"文献1\" or any number prefix\n"
                            knowledge_base_context += f"3. **ABSOLUTELY FORBIDDEN**: Using PDF content snippets as document names (e.g., \"Studies in Language Companion Series\" or \"(Studies in Language Companion Series) Matti M\" is WRONG)\n"
                            knowledge_base_context += f"4. **REQUIRED**: Use the exact names from the Document Name List, do NOT modify, truncate, or extract from content\n"
                            knowledge_base_context += f"5. Correct example: \"According to {list(document_mapping.values())[0] if document_mapping else '[Document Name]'}\" ✅\n"
                            knowledge_base_context += f"6. Wrong example: \"According to Document 1: {list(document_mapping.values())[0] if document_mapping else '[Document Name]'}\" ❌ (do not use number prefix)\n"
                            knowledge_base_context += f"7. Wrong example: \"Document 1: Studies in Language Companion Series\" ❌ (this is a content snippet, not a document name)\n"
                    
                    # 第三步：展示文档内容（明确标注文档名称，避免混淆）
                    for index, result in enumerate(kb_results[:10], 1):
                        document = result.get("document", {})
                        content = result.get("content") or document.get("content", "")
                        source_name = document_mapping.get(index, f"文档 {index}" if user_lang == "zh" else f"Document {index}")
                        
                        doc_label = "文献" if user_lang == "zh" else "Document"
                        content_label = "内容" if user_lang == "zh" else "Content"
                        name_label = "文档名称" if user_lang == "zh" else "Document Name"
                        
                        # 明确标注：这是文档名称，不是内容
                        knowledge_base_context += f"\n**{doc_label} {index}**\n"
                        knowledge_base_context += f"**{name_label}**: {source_name}\n"
                        knowledge_base_context += f"**{content_label}**: {content}\n"
                        knowledge_base_context += "\n---\n"
                    
                    # 打印文档映射信息到日志
                    logger.info("  - 文档映射:")
                    for doc_index, doc_name in document_mapping.items():
                        logger.info(f"    文献{doc_index}: {doc_name}")
                    
                    # 不再从知识库内容中提取特征ID，因为：
                    # 1. PDF原文中通常不包含Grambank/D-PLACE的特征ID（如GB051, GB054等）
                    # 2. 如果提取，可能会误导LLM认为这些ID是文献中提到的
                    # 3. 特征ID应该从数据库搜索结果中获取，而不是从PDF文献中提取
                    # extracted_features = extract_features_from_knowledge_base(kb_results)
                    # if extracted_features:
                    #     features_label = "从文献中提取的特征ID" if user_lang == "zh" else "Feature IDs Extracted from Literature"
                    #     knowledge_base_context += f"\n**{features_label}**: {', '.join(extracted_features[:15])}\n"
                    # logger.info(f"  - 从知识库提取的特征ID: {extracted_features[:10]}")
            except Exception as e:
                logger.warning(f"知识库搜索失败（不影响推荐）: {e}")
        else:
            logger.info("  - 知识库未初始化，跳过搜索")
        
        # 4. 构建 LLM 提示
        logger.info("📝 步骤4: 构建LLM提示词")
        # 构建特征列表字符串（自然语言格式，不强制排名）
        category_label = "分类" if user_lang == "zh" else "Category"
        desc_label = "描述" if user_lang == "zh" else "Description"
        features_text = "\n".join([
            f"{feature.get('id', '')} ({feature.get('source', '')}): {feature.get('name', '')}\n  {category_label}: {feature.get('category', '')}\n  {desc_label}: {clean_description(feature.get('description', ''))}"
            for feature in search_results
        ])
        
        # 根据语言构建不同的提示词
        if user_lang == "zh":
            prompt = f"""你是一位专业的语言学数据分析专家。现在你有机会探索完整的语言学数据库和知识库（研究文献）来为用户推荐最相关的特征。

=== 用户问题 ===
{request.user_query}

=== 完整数据库信息 ===
Grambank数据库: {stats.get('totalGrambankFeatures', 0)} 个语法特征
D-PLACE数据库: {stats.get('totalDplaceFeatures', 0)} 个社会文化特征
WALS数据库: {stats.get('totalWalsFeatures', 0)} 个语言结构特征

=== 数据库中的相关特征（共找到{len(search_results)}个） ===
以下是从数据库中搜索到的可能与你的查询相关的特征。请仔细阅读每个特征的名称和描述，根据你的查询意图选择最相关的特征。

**重要**：你只能从下面列出的特征中选择特征ID，不能自己编造特征ID。如果某个特征不在下面的列表中，即使它看起来相关，也不能使用。

{features_text}
{knowledge_base_context}

=== 任务 ===
请仔细分析用户的问题，理解用户的查询意图，然后从以下来源推荐最相关的特征：
1. **必须**：从上面搜索到的相关特征列表中选择特征ID。**绝对不能自己编造特征ID**，只能使用上面列表中实际存在的特征ID
2. 如果知识库（PDF文献）中有相关研究，可以参考，但这不是必须的
3. 如果上面搜索到的特征列表中没有足够相关的特征，**只能从列表中选择最相关的特征，即使数量较少也没关系**。不能编造不存在的特征ID

**重要**：
- **所有特征ID必须来自上面提供的特征列表，不能自己编造**
- 请根据特征的实际内容和你的查询意图来判断相关性，而不是简单地按顺序选择
- 选择那些真正能回答用户问题的特征
- 如果列表中没有完全匹配的特征，选择最接近的即可，但必须是列表中实际存在的特征ID

**关于source字段**：
- source字段必须填写特征的**实际来源数据库**，可以是"Grambank"、"D-PLACE"或"WALS"
- 如果特征来自知识库文献的引用，source填写"Knowledge Base"

请按以下JSON格式返回推荐（必须使用中文）：

{{
  "recommendations": [
    {{
      "category": "特征分类名称",
      "name": "推荐组名称", 
      "description": "推荐理由",
      "features": ["特征ID", "特征ID", "特征ID"],
      "reason": "为什么推荐这些特征（主要基于特征的实际内容和相关性。如果知识库中有相关研究，可以引用，但这不是必须的）。**重要**：如果引用知识库，直接使用文档名称，格式如\"根据[文档名称]\"，不要使用\"文献1\"、\"Document 1\"等编号前缀，不要使用PDF内容片段作为文档名称",
      "source": "Grambank"、"D-PLACE"或"WALS"（根据特征ID判断，GB开头=Grambank，EA开头或其他=D-PLACE，数字+字母开头如1A=WALS）
    }}
  ]
}}

**重要：特征ID格式要求**：
- features字段中的特征ID必须是纯ID格式，例如："GB051"、"EA046"、"52A"、"1A"等
- **不要**在特征ID中添加后缀，例如：不要写成"52A (WALS)"或"GB051 (Grambank)"，只写"52A"或"GB051"即可
- source字段会单独标注特征的来源数据库，不需要在特征ID中重复

要求：
1. 仔细理解用户的查询意图，选择真正能回答用户问题的特征
2. **必须**从上面搜索到的特征列表中选择特征ID，不能自己编造特征ID。如果某个特征不在列表中，即使看起来相关也不能使用
3. 从搜索到的特征列表中选择最相关的特征（不一定要按顺序，而是根据实际相关性）
4. 如果知识库中有相关研究，可以参考并引用，但这不是必须的。如果知识库中没有相关内容，直接基于特征列表推荐即可
5. **绝对禁止编造特征ID**。如果列表中没有足够相关的特征，就选择最相关的几个即可，即使数量较少也没关系。不能因为想要推荐更多特征而编造不存在的特征ID
6. 考虑特征之间的关联性和互补性
7. 提供清晰的推荐理由，如果知识库中有相关研究可以结合，但主要基于特征的实际内容和相关性
8. 标注特征来源（Grambank/D-PLACE/WALS/Knowledge Base）
9. 所有文本内容必须使用中文
8. **引用格式要求（必须严格遵守）**：
   - 引用知识库文献时，直接使用文档名称，格式示例："根据[文档名称]" 或 "[文档名称]提到..."
   - **禁止**：使用"文献1"、"Document 1"等编号前缀
   - **禁止**：使用PDF内容片段作为文档名称（如"Studies in Language Companion Series"是错误的）
   - **必须**：使用文档名称列表中列出的确切名称，不要修改或截断"""
        else:
            prompt = f"""You are a professional linguistic data analysis expert. You now have the opportunity to explore the complete linguistic database and knowledge base (research literature) to recommend the most relevant features for the user.

=== User Query ===
{request.user_query}

=== Complete Database Information ===
Grambank Database: {stats.get('totalGrambankFeatures', 0)} grammatical features
D-PLACE Database: {stats.get('totalDplaceFeatures', 0)} social-cultural features
WALS Database: {stats.get('totalWalsFeatures', 0)} language structure features

=== Related Features Found in Database ({len(search_results)} features found) ===
The following features were found in the database that may be relevant to your query. Please carefully read each feature's name and description, and select the most relevant features based on your query intent.

**IMPORTANT**: You can ONLY select feature IDs from the list below. Do NOT make up feature IDs. If a feature is not in the list below, you cannot use it even if it seems relevant.

{features_text}
{knowledge_base_context}

=== Task ===
Please carefully analyze the user's query, understand the query intent, and then recommend the most relevant features from the following sources:
1. **MUST**: Select feature IDs from the related features list above. **DO NOT make up feature IDs**. You can ONLY use feature IDs that actually exist in the list above
2. If there is relevant research in the knowledge base (PDF literature), you can reference it, but this is optional
3. If there are not enough relevant features in the list above, **you can only select the most relevant features from the list, even if the number is small**. Do NOT make up non-existent feature IDs

**Important**:
- **All feature IDs MUST come from the feature list provided above. Do NOT make up feature IDs**
- Please judge relevance based on the actual content of features and your understanding of the query intent, rather than simply selecting in order
- Choose features that can truly answer the user's question
- If there are no perfectly matching features in the list, select the closest ones, but they MUST be actual feature IDs from the list

**About the source field**:
- The source field must indicate the **actual source database** of the feature, which can be "Grambank", "D-PLACE", or "WALS"
- If a feature is cited from knowledge base literature, use "Knowledge Base" as source

Please return recommendations in the following JSON format (must use English):

{{
  "recommendations": [
    {{
      "category": "Feature Category Name",
      "name": "Recommendation Group Name", 
      "description": "Recommendation Reason",
      "features": ["FeatureID", "FeatureID", "FeatureID"],
      "reason": "Why these features are recommended (primarily based on the actual content and relevance of features. If there is relevant research in the knowledge base, you can cite it, but this is optional). **IMPORTANT**: If citing the knowledge base, directly use the document name, format like \"According to [Document Name]\", do NOT use \"Document 1\" or any number prefix, do NOT use PDF content snippets as document names",
      "source": "Grambank", "D-PLACE", or "WALS" (determined by feature ID: GB prefix = Grambank, EA prefix or others = D-PLACE, number+letter like 1A = WALS)
    }}
  ]
}}

**IMPORTANT: Feature ID Format Requirements**:
- Feature IDs in the features field must be in pure ID format, e.g., "GB051", "EA046", "52A", "1A", etc.
- **DO NOT** add suffixes to feature IDs, e.g., do NOT write "52A (WALS)" or "GB051 (Grambank)", just write "52A" or "GB051"
- The source field will separately label the feature's source database, no need to repeat it in the feature ID

Requirements:
1. Carefully understand the user's query intent and select features that can truly answer the user's question
2. **MUST** select feature IDs from the search results list above. Do NOT make up feature IDs. If a feature is not in the list, you cannot use it even if it seems relevant
3. Select the most relevant features from the search results list (not necessarily in order, but based on actual relevance)
4. If there is relevant research in the knowledge base, you can reference and cite it, but this is optional. If there is no relevant content in the knowledge base, simply recommend based on the feature list
5. **ABSOLUTELY FORBIDDEN to make up feature IDs**. If there are not enough relevant features in the list, just select the most relevant ones, even if the number is small. Do NOT make up non-existent feature IDs just because you want to recommend more features
6. Consider relationships and complementarity between features
7. Provide clear recommendation reasons, primarily based on the actual content and relevance of features. If there is relevant research in the knowledge base, you can combine it, but it's not required
8. Label feature sources (must be "Grambank", "D-PLACE", "WALS", or "Knowledge Base" based on feature ID, not "Current Data")
9. All text content must be in English
8. **Citation Format Requirements (MUST strictly follow)**:
   - When citing knowledge base documents, directly use the document name, format example: "According to [Document Name]" or "[Document Name] mentions..."
   - **FORBIDDEN**: Using "Document 1", "文献1" or any number prefix
   - **FORBIDDEN**: Using PDF content snippets as document names (e.g., "Studies in Language Companion Series" is WRONG)
   - **REQUIRED**: Use the exact names from the Document Name List, do NOT modify or truncate them"""
        
        logger.info(f"  - 提示词长度: {len(prompt)} 字符")
        logger.info(f"  - 提示词完整内容:\n{prompt}")
        
        # 5. 调用 LLM 获取推荐
        logger.info("🤖 步骤5: 调用LLM API获取推荐")
        logger.info(f"  - 使用API: {api_provider}")
        if api_provider == "qianwen":
            response_text = await call_qianwen_api_for_recommendation(prompt, api_key)
        else:
            response_text = await call_gemini_api(prompt, api_key)
        
        logger.info(f"  - LLM响应长度: {len(response_text)} 字符")
        logger.info(f"  - LLM响应完整内容:\n{response_text}")
        
        # 6. 解析响应
        logger.info("🔍 步骤6: 解析LLM响应")
        import re
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            try:
                parsed = json.loads(json_match.group(0))
                logger.info(f"  - 解析成功，找到 {len(parsed.get('recommendations', []))} 个推荐")
                
                if parsed.get("recommendations"):
                    # 验证特征是否存在
                    logger.info("  - 验证特征有效性")
                    all_feature_ids = get_all_feature_ids()
                    valid_gb = set(all_feature_ids.get('grambank', []))
                    valid_ea = set(all_feature_ids.get('dplace', []))
                    valid_wals = set(all_feature_ids.get('wals', []))
                    
                    for i, rec in enumerate(parsed["recommendations"]):
                        original_count = len(rec.get("features", []))
                        valid_features = []
                        for feature_id in rec.get("features", []):
                            # 清理特征ID：移除可能的后缀如 "(WALS)", "(Grambank)", "(D-PLACE)" 等
                            import re
                            cleaned_id = re.sub(r'\s*\(WALS\)\s*$', '', feature_id, flags=re.IGNORECASE)
                            cleaned_id = re.sub(r'\s*\(Grambank\)\s*$', '', cleaned_id, flags=re.IGNORECASE)
                            cleaned_id = re.sub(r'\s*\(D-PLACE\)\s*$', '', cleaned_id, flags=re.IGNORECASE)
                            cleaned_id = cleaned_id.strip()
                            
                            # 判断特征属于哪个数据库
                            if cleaned_id in valid_gb:
                                valid_features.append({'id': cleaned_id, 'source': 'Grambank'})
                            elif cleaned_id in valid_ea:
                                valid_features.append({'id': cleaned_id, 'source': 'D-PLACE'})
                            elif cleaned_id in valid_wals:
                                valid_features.append({'id': cleaned_id, 'source': 'WALS'})
                            else:
                                # 检查是否是D-PLACE的其他格式（CARNEIRO_, B前缀, SCCS, Richness等）
                                is_dplace_other = (
                                    cleaned_id.startswith('CARNEIRO_') or
                                    cleaned_id.startswith('SCCS') or
                                    'Richness' in cleaned_id or
                                    (cleaned_id.startswith('B') and not cleaned_id.startswith('GB') and re.match(r'^B\d{1,4}$', cleaned_id)) or
                                    re.match(r'^(Annual|Monthly|Net|Precipitation|Temperature|Biome|EcoRegion|Elevation|Slope|DistToCoast)', cleaned_id)
                                )
                                if is_dplace_other and cleaned_id in valid_ea:
                                    # 这些特征在valid_ea中（因为load_dplace_variables会加载所有D-PLACE变量，包括CARNEIRO_, B等）
                                    valid_features.append({'id': cleaned_id, 'source': 'D-PLACE'})
                                else:
                                    logger.warning(f"    - 无效特征ID: {feature_id} (清理后: {cleaned_id}, 不在任何数据库中，LLM可能编造了这个ID)")
                        rec["features"] = valid_features
                        logger.info(f"    - 推荐 {i+1} ({rec.get('name', 'N/A')}): {original_count} -> {len(valid_features)} 个有效特征")
                    
                    # 过滤掉没有有效特征的推荐
                    before_filter = len(parsed["recommendations"])
                    parsed["recommendations"] = [
                        rec for rec in parsed["recommendations"]
                        if len(rec.get("features", [])) > 0
                    ]
                    after_filter = len(parsed["recommendations"])
                    logger.info(f"  - 过滤后: {before_filter} -> {after_filter} 个推荐")
                    
                    logger.info("✅ 特征推荐流程完成")
                    logger.info("=" * 60)
                    return parsed
            except json.JSONDecodeError as e:
                logger.warning(f"  - LLM响应解析失败: {e}")
                logger.warning(f"  - 响应完整内容:\n{response_text}")
        
        # 如果解析失败，返回空推荐
        logger.warning("⚠️ 解析失败，返回空推荐")
        logger.info("=" * 60)
        return {"recommendations": []}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"特征推荐失败: {e}")
        raise HTTPException(status_code=500, detail=f"特征推荐失败: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=HOST, port=PORT) 